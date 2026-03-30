"""
DOCX 내보내기 - 마크다운 아티팩트를 DOCX 파일로 변환
"""

from __future__ import annotations

from pathlib import Path
from typing import Optional

from .logger import get_logger

logger = get_logger(__name__)

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None


def is_docx_available() -> bool:
    """python-docx 패키지 사용 가능 여부"""
    return Document is not None


def export_to_docx(
    export_md_path: Path,
    output_path: Path,
    project_info: dict | None = None,
) -> Optional[Path]:
    """
    마크다운 내보내기 파일을 DOCX로 변환합니다.

    Args:
        export_md_path: export.md 파일 경로
        output_path: 출력 DOCX 파일 경로
        project_info: 프로젝트 메타데이터 (회사명, 직무 등)

    Returns:
        생성된 DOCX 파일 경로 또는 None (실패 시)
    """
    if not is_docx_available():
        logger.warning(
            "python-docx 패키지가 설치되어 있지 않아 DOCX 내보내기를 건너뜁니다."
        )
        return None

    if not export_md_path.exists():
        logger.error(f"내보내기 파일을 찾을 수 없습니다: {export_md_path}")
        return None

    try:
        md_text = export_md_path.read_text(encoding="utf-8")
        doc = Document()

        _configure_styles(doc)

        if project_info:
            _add_title_page(doc, project_info)

        _convert_markdown_to_docx(doc, md_text)

        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        logger.info(f"DOCX 내보내기 완료: {output_path}")
        return output_path

    except Exception as e:
        logger.error(f"DOCX 내보내기 실패: {e}")
        return None


def _configure_styles(doc: "Document") -> None:
    """문서 스타일 설정"""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "맑은 고딕"
    font.size = Pt(11)

    for level in range(1, 4):
        heading_style = doc.styles[f"Heading {level}"]
        heading_style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)


def _add_title_page(doc: "Document", project_info: dict) -> None:
    """표지 페이지 추가"""
    company = project_info.get("company_name", "") or "지원 회사"
    role = project_info.get("job_title", "") or "지원 직무"

    title = doc.add_heading(f"{company} - {role}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")
    subtitle = doc.add_paragraph("취업 준비 패키지")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()


def _convert_markdown_to_docx(doc: "Document", md_text: str) -> None:
    """마크다운 텍스트를 DOCX 구조로 변환"""
    lines = md_text.split("\n")
    in_code_block = False
    code_buffer: list[str] = []

    for line in lines:
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code_block:
                _add_code_block(doc, "\n".join(code_buffer))
                code_buffer = []
            in_code_block = not in_code_block
            continue

        if in_code_block:
            code_buffer.append(line)
            continue

        if not stripped:
            continue

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith("#### "):
            doc.add_heading(stripped[5:].strip(), level=4)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
        elif stripped[0:2].isdigit() and len(stripped) > 2 and stripped[2] == ".":
            doc.add_paragraph(stripped[3:].strip(), style="List Number")
        elif stripped.startswith(">"):
            quote_text = stripped.lstrip("> ").strip()
            p = doc.add_paragraph(quote_text)
            p.runs[0].italic = True
            p.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        else:
            p = doc.add_paragraph(stripped)
            _apply_inline_formatting(p, stripped)

    if in_code_block and code_buffer:
        _add_code_block(doc, "\n".join(code_buffer))


def _add_code_block(doc: "Document", code: str) -> None:
    """코드 블록 추가"""
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def _apply_inline_formatting(paragraph, text: str) -> None:
    """인라인 마크다운 서식 적용 (굵게, 기울임 등)"""
    import re

    paragraph.clear()

    pattern = r"(\*\*\*[^*]+\*\*\*|\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)"
    parts = re.split(pattern, text)

    for part in parts:
        if not part:
            continue
        if part.startswith("***") and part.endswith("***"):
            run = paragraph.add_run(part[3:-3])
            run.bold = True
            run.italic = True
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        else:
            paragraph.add_run(part)


def export_artifacts_to_docx(
    coach_path: Path,
    writer_path: Path,
    interview_path: Path,
    output_path: Path,
    project_info: dict | None = None,
) -> Optional[Path]:
    """
    개별 아티팩트 파일들을 하나의 DOCX로 통합합니다.

    Args:
        coach_path: coach.md 경로
        writer_path: writer.md 경로
        interview_path: interview.md 경로
        output_path: 출력 DOCX 경로
        project_info: 프로젝트 메타데이터

    Returns:
        생성된 DOCX 파일 경로 또는 None
    """
    if not is_docx_available():
        logger.warning(
            "python-docx 패키지가 설치되어 있지 않아 DOCX 내보내기를 건너뜁니다."
        )
        return None

    try:
        doc = Document()
        _configure_styles(doc)

        if project_info:
            _add_title_page(doc, project_info)

        sections = [
            ("Coach - 코칭 결과", coach_path),
            ("Writer - 자소서 작성 결과", writer_path),
            ("Interview - 면접 준비 결과", interview_path),
        ]

        for title, path in sections:
            doc.add_page_break()
            doc.add_heading(title, level=1)
            if path.exists():
                content = path.read_text(encoding="utf-8")
                _convert_markdown_to_docx(doc, content)
            else:
                p = doc.add_paragraph("(아티팩트가 생성되지 않았습니다)")
                p.runs[0].italic = True
                p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        logger.info(f"통합 DOCX 내보내기 완료: {output_path}")
        return output_path

    except Exception as e:
        logger.error(f"통합 DOCX 내보내기 실패: {e}")
        return None
